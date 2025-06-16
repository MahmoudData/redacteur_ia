import re
from collections import defaultdict
import json
from langchain_community.document_loaders.llmsherpa import LLMSherpaFileLoader
from collections import defaultdict
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openai import OpenAI
from typing import List
from qdrant_client import QdrantClient, models
from datetime import datetime
import uuid
from docx import Document
from io import BytesIO
import time


# ----- structuration du sommaire -----
def parser_sommaire(sommaire: str) -> dict:

    regex_sous_sous_section = re.compile(r'^\s*(\d+\.\d+\.\d+)\.\s+(.*)', re.MULTILINE)
    regex_sous_section = re.compile(r'^\s*(\d+\.\d+)\.\s+(.*)', re.MULTILINE)
    regex_section = re.compile(r'^\s*(\d+)\.\s+(.*)', re.MULTILINE)

    structure = defaultdict(lambda: {
        "title": "",
        "sous_sections": defaultdict(lambda: {
            "title": "",
            "sous_sous_sections": {}
        })
    })

    lines = sommaire.strip().split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if (m := regex_sous_sous_section.match(line)):
            ref, title = m.groups()
            parts = ref.split(".")
            section = parts[0]
            sous_section = f"{parts[0]}.{parts[1]}"
            structure[section]["sous_sections"][sous_section]["sous_sous_sections"][ref] = title

        elif (m := regex_sous_section.match(line)):
            ref, title = m.groups()
            parts = ref.split(".")
            section = parts[0]
            structure[section]["sous_sections"][ref]["title"] = title

        elif (m := regex_section.match(line)):
            ref, title = m.groups()
            structure[ref]["title"] = title

    return json.loads(json.dumps(structure))


# ----- génération des prompts -----
def generate_prompts_from_structure(structure: dict) -> list[dict]:
    prompts = []
    titres_utilises = []

    def parcourir_niveaux(niveau: int, titre: str, enfants: dict, parent_titre: str = None):
        for ref, contenu in enfants.items():
            if niveau == 1:
                titre_complet = f"{ref}. {contenu['title']}"
                sous_sections = contenu.get("sous_sections", {})
                if not sous_sections:
                    prompts.append({
                        "prompt": f"Rédige la section : {titre_complet}. Ne reprend pas le titre dans ta réponse"
                    })
                    titres_utilises.append(titre_complet)
                else:
                    parcourir_niveaux(2, titre_complet, sous_sections, titre_complet)

            elif niveau == 2:
                titre_complet = f"{ref}. {contenu['title']}"
                sous_sous_sections = contenu.get("sous_sous_sections", {})
                if not sous_sous_sections:
                    prompts.append({
                        "prompt": f"Rédige la section : {titre_complet} issue de la section : {parent_titre}. Ne reprend pas le titre dans ta réponse"
                    })
                    titres_utilises.append(titre_complet)
                else:
                    parcourir_niveaux(3, titre_complet, sous_sous_sections, titre_complet)

            elif niveau == 3:
                titre_complet = f"{ref}. {contenu}"
                prompts.append({
                    "prompt": f"Rédige la section : {titre_complet} issue de la section : {parent_titre}. Ne reprend pas le titre dans ta réponse"
                })
                titres_utilises.append(titre_complet)

    parcourir_niveaux(1, None, structure)

    return [
        {
            "prompt": prompt["prompt"],
            "titre": titres_utilises[i]
        }
        for i, prompt in enumerate(prompts)
    ]


# ----- chargement des documents sous forme structurée -----
def load_documents_from_paths(
    pdf_paths: list,
    llmsherpa_api_url: str,
    apply_ocr: bool = False,
    strategy: str = "sections"
) -> list:
    """
    Charge des documents à partir de chemins PDF via LLM Sherpa.
    Retourne une liste de documents LangChain.
    """
    all_docs = []

    for path in pdf_paths:
        loader = LLMSherpaFileLoader(
            file_path=path,
            new_indent_parser=True,
            apply_ocr=apply_ocr,
            strategy=strategy,
            llmsherpa_api_url=llmsherpa_api_url,
        )
        docs = loader.load()
        all_docs.extend(docs)

    return all_docs


# ----- groupement des documents par section -----
def group_docs_by_section(docs):
    section_map = defaultdict(list)

    for doc in docs:
        section = doc.metadata.get("section_title", "Sans titre")
        section_map[section].append(doc.page_content)

    return section_map


# ----- decoupage du texte -----
def split_text(text: list[str], chunk_size: int = 1000, chunk_overlap: int = 100) -> List[str]:

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap = chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""],
    )

    all_chunks = []
    for txt in text:
        chuncks = text_splitter.split_text(txt)
        all_chunks.extend(chuncks)
        
    return all_chunks


# ----- decoupage des sections avec contexte -----
def chunk_sections_with_context(section_map: dict, chunk_size=1000, chunk_overlap=100):
    all_chunks = []

    for section_title, paragraphs in section_map.items():
        full_text = "\n".join(paragraphs)  
        chunks = split_text([full_text], chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        for chunk in chunks:
            all_chunks.append({
                "section": section_title,
                "text": chunk
            })

    return all_chunks


# ----- embedding des chunks -----
def embed_text(chunks: List, model: str, api_key: str, batch_size: int = 20, sleep_time: float = 1.0):
    # Extraire les textes
    if isinstance(chunks[0], str):
        texts = chunks
    else:
        texts = [c["text"] for c in chunks]

    client = OpenAI(
        base_url="https://api.scaleway.ai/v1",
        api_key=api_key
    )

    results = []

    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]

        try:
            response = client.embeddings.create(
                input=batch,
                model=model
            )
            results.extend([
                {"text": batch[j], "embedding": d.embedding}
                for j, d in enumerate(response.data)
            ])
        except Exception as e:
            print(f"Erreur sur le batch {i // batch_size + 1}: {e}")
            time.sleep(5)  # pause plus longue en cas d'erreur
            continue

        time.sleep(sleep_time)  # petite pause pour éviter d'être rate-limité

    return results



# ----- creation de la collection Qdrant -----
def create_collection(host: str, vector_size: int = 100):

    now = datetime.now()
    collection_name  =f"session_{int(now.timestamp() * 1000)}"

    client = QdrantClient(url=host)

    client.create_collection(
        collection_name = collection_name,
        vectors_config=models.VectorParams(size=vector_size, distance=models.Distance.COSINE),
    )

    return collection_name


# ----- insertion des chunks dans Qdrant avec batching -----
def upsert_embeddings(embeddings, collection_name: str, host: str, batch_size: int = 100):
    client = QdrantClient(url=host)

    for i in range(0, len(embeddings), batch_size):
        batch = embeddings[i:i + batch_size]

        points = [
            models.PointStruct(
                id=str(uuid.uuid4()),
                payload={
                    "text": doc["text"],
                    "section": doc.get("section", "")
                },
                vector=doc["embedding"],
            )
            for doc in batch
        ]

        client.upsert(
            collection_name=collection_name,
            points=points
        )


# ----- recherche dans Qdrant -----
def search_in_qdrant(
    query_text: str,
    embedding_fn,
    model: str,
    api_key: str,
    collection_name: str,
    host: str,
    top_k: int = 1
) -> List[str]:
    """
    Effectue une recherche vectorielle dans Qdrant à partir d'une question.
    Retourne les meilleurs chunks de texte associés.
    """
    embedded = embedding_fn([query_text], model=model, api_key=api_key)[0]["embedding"]

    client = QdrantClient(url=host)

    results = client.search(
        collection_name=collection_name,
        query_vector=embedded,
        limit=top_k
    )

    return [r.payload["text"] for r in results]


# ----- configuration Scaleway LLM (Llama 3) -----
def scaleway_llm(system: str, user: str, model: str, api_key: str) -> str:

  client = OpenAI(
      base_url = "https://api.scaleway.ai/26a51179-e2ad-495f-b452-03f48b7a2434/v1",
      api_key = api_key
  )

  response = client.chat.completions.create(
      model="llama-3.3-70b-instruct",
      messages=[
          { "role": "system", "content": system },
      { "role": "user", "content": user },
      ],
      max_tokens=1024,
      temperature=0.6,
      top_p=0.9,
      presence_penalty=0,
      stream=False,
  )
  return response.choices[0].message.content


# ----- génération de la réponse avec OpenAI LLM -----
def generate_sections_from_prompts(
    prompts: list,
    titre_doc: str,
    finalite_doc: str,
    collection_name: str,
    embedding_fn,
    llm_fn,
    model_embed: str,
    model_llm: str,
    api_key: str,
    host: str,
    top_k: int = 5
):
    results = []

    for p in prompts:
        prompt_text = p["prompt"]
        titre_section = p["titre"]

        chunks = search_in_qdrant(
            query_text=prompt_text,
            embedding_fn=embedding_fn,
            model=model_embed,
            api_key=api_key,
            collection_name=collection_name,
            host=host,
            top_k=top_k
        )

        context = "\n".join(chunks)
        prompt_system = f"""Tu rédige un document intitulé **{titre_doc}** destiné à **{finalite_doc}**.

Consignes :
Tu dois reprendre l'intégralité des données techniques, valeurs chiffrées, hypothèses de calcul, normes, configurations et paramètres fournis dans le contexte.
Adapte le ton pour qu’il soit conforme à un manuel technique clair : phrases structurées, style informatif, voix passive autorisée.
Ne fais aucune interprétation personnelle, ne complète pas avec des données externes.
Tu peux reformuler légèrement pour améliorer la lisibilité, mais aucune information ne doit être omise ou modifiée.

----------------
Context: 
{context}
"""

        answer = llm_fn(
            system=prompt_system,
            user=prompt_text,
            model=model_llm,
            api_key=api_key
        )

        results.append({
            "titre": titre_section,
            "contenu": answer
        })

        time.sleep(5)

    return results


# ----- suppression de la collection Qdrant -----
def delete_collection(collection_name: str, host: str):
    """
    Supprime une collection dans Qdrant.
    """
    client = QdrantClient(url=host)

    client.delete_collection(collection_name=collection_name)


# --- Génération du .docx ---
def generer_docx(result_json):
    doc = Document()

    doc.add_heading(result_json["titre_document"], level=0)

    doc.add_heading("Sommaire", level=1)
    for ligne in result_json.get("sommaire", []):
        doc.add_paragraph(ligne, style="List Bullet")

    doc.add_heading("Contenu", level=1)
    for section in result_json.get("sections", []):
        doc.add_heading(section["titre"], level=2)
        doc.add_paragraph(section["contenu"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

